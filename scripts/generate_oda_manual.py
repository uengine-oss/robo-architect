"""043 ODA 표준 분해 모드 — DOCX 사용자 매뉴얼 생성기.

frontend/test-results/demo-*.png 스크린샷을 엔드유저 관점 DOCX 매뉴얼로 엮는다.
run: uv run --with python-docx python scripts/generate_oda_manual.py
"""

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).resolve().parents[1]
SHOTS = BASE / "frontend" / "test-results"
OUT = BASE / "specs" / "043-oda-standard-mode" / "manual.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

ACCENT = RGBColor(0x1A, 0x73, 0xE8)
GRAY = RGBColor(0x71, 0x80, 0x96)


def add_screenshot(doc, filename, caption, width=6.2):
    path = SHOTS / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[이미지 없음: {filename}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0, 0)
    cap = doc.add_paragraph(f"▲ {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GRAY


def add_note(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(f"💡 {text}"); r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)


def add_warning(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(f"⚠️  {text}"); r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0xC0, 0x53, 0x21)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.4); s.bottom_margin = Cm(2.4)
    s.left_margin = Cm(2.8); s.right_margin = Cm(2.4)

# 표지
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ODA 표준 분해 모드\n사용자 매뉴얼")
r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = ACCENT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Robo Architect · Proposals — 043"); r.font.size = Pt(13); r.font.color.rgb = GRAY
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"발행일 {date.today().isoformat()}   ·   대상: 아키텍트 / 설계자")
r.font.size = Pt(10); r.font.color.rgb = GRAY
doc.add_page_break()

# 1. 소개
h(doc, "1. 서비스 소개")
doc.add_paragraph(
    "ODA 표준 분해 모드는 Proposal(설계 변경 제안)을 TM Forum ODA 표준에 근거해 분해·검증·"
    "설계하는 모드입니다. 요청을 표준에 매핑하고, 표준이 이미 정의한 것을 재사용하며, 모든 "
    "요소를 REUSE/EXTEND/NEW 로 분류하고, ‘준수 후 확장(comply-then-extend)’을 차단형 "
    "적합성 게이트로 강제합니다."
)
for b in [
    "분해 모드 3종: 간소화 · 상세 DDD · ODA 표준(신규)",
    "표준 정합성 매핑: Use Case(UCxxx)·SID 엔티티·TMF Open API·ODA Component 블록",
    "적합성 게이트: REUSE/EXTEND/NEW 분류 + 표준 위반 시 진행 차단(면제 가능)",
    "표준 산출물: SID 데이터 모델 · TMF 계약 · ODA 아키텍처 · BDD .feature",
]:
    doc.add_paragraph(b, style="List Bullet")

# 2. 시작하기 전에
h(doc, "2. 시작하기 전에")
doc.add_paragraph("Proposals 탭에서 새 Proposal 을 만들 때 분해 모드를 선택합니다. ODA 표준 모드는 "
                  "ODA 표준 지식 베이스(SID·use-case 라이브러리)가 있어야 동작합니다.")
add_warning(doc, "지식 베이스를 찾지 못하면 ODA 분해가 ‘지식 베이스를 찾을 수 없습니다’로 중단됩니다.")

STEPS = [
    ("STEP 1. Proposals 화면", "demo-01-proposals-list.png",
     "Proposals 탭. 오른쪽 위 ‘＋ 새 Proposal’ 로 제안 작성을 시작합니다.",
     "각 Proposal 은 분해 모드와 상태(INTENT/PLAN/…)로 구분됩니다."),
    ("STEP 2. 분해 모드 선택 — 3종", "demo-02-create-three-modes.png",
     "요구사항을 자연어로 입력하고, 아래 모드 스위치에서 간소화 · 상세 DDD · ODA 표준 중 하나를 고릅니다.",
     "ODA 표준은 TM Forum 표준 근거 설계·검증을 수행하는 신규 옵션입니다."),
    ("STEP 3. ODA 표준 모드로 분석 시작", "demo-03-create-oda-selected.png",
     "‘ODA 표준’을 선택하면 강조 표시됩니다. ‘AI 분석 시작’을 누르면 Proposal 이 생성되고 상세 화면의 "
     "Intent 탭에 ODA 트랙이 열립니다.",
     "표준 정합성 분해는 ODA 지식 베이스를 근거로 진행됩니다."),
    ("STEP 4. 표준 정합성 매핑", "demo-04-oda-alignment.png",
     "① 표준 정합성 매핑 영역에서 매칭된 Use Case(UC003), 재사용할 SID 엔티티(Customer·ProductOrder), "
     "기준 TMF API(TMF622 v4), 대상 Component 블록(coreFunction)을 확인합니다.",
     "표준 정렬된 strategicDiff 도 함께 산출되어 Impact 탭에서 볼 수 있습니다."),
    ("STEP 5. 적합성 게이트 — 분류와 위반", "demo-05-oda-conformance-fail.png",
     "② 적합성 게이트에서 각 요소가 REUSE/EXTEND/NEW 로 분류됩니다. 표준을 깨는 위반이 있으면 게이트가 "
     "‘FAIL — 진행 차단’ 으로 표시되어 plan 수립·제출이 막힙니다.",
     "예시 위반: ProductOrder.orderDate 의 표준 필드 타입 재정의(재타이핑) — 추가형이 아님."),
    ("STEP 6. 위반 면제 사유 입력", "demo-06-oda-waive-reason.png",
     "FAIL 을 진행하려면 위반 목록 아래 입력란에 ‘면제 사유’를 적습니다. 사유는 필수이며 적합성 "
     "리포트에 기록됩니다.",
     "면제는 표준을 깨는 변경을 의도적으로 수용할 때만 사용하세요."),
    ("STEP 7. 면제 후 — 게이트 WAIVED", "demo-07-oda-gate-waived.png",
     "‘면제하고 진행’을 누르면 게이트가 ‘면제됨(WAIVED)’ 이 되고, ③ 표준 산출물의 ‘표준 설계(plan) "
     "실행’ 버튼이 활성화됩니다.",
     "사유 없이 제출/plan 을 시도하면 서버가 차단합니다(면제만이 통과 경로)."),
    ("STEP 8. 표준 설계 산출물", "demo-08-oda-artifacts.png",
     "③ 표준 산출물에서 SID 데이터 모델 · TMF 계약 · ODA 아키텍처 · BDD .feature 네 가지 산출물을 "
     "생성·확인합니다. 결과는 표준 tacticalDiff 로 수렴되어 이후 단계가 분기 없이 진행됩니다.",
     "‘Plan 단계로 진행’을 누르면 제출(DRAFT→SUBMITTED)되며, 게이트가 차단 상태면 이때도 막힙니다."),
    ("STEP 9. Plan 단계 (Constitution 기반 구현계획)", "demo-09-oda-plan-stage.png",
     "제출 후 Plan 탭에서는 041 Constitution 기반 구현계획이 그대로 표시됩니다 — 배포 환경"
     "(Kubernetes + ODA Canvas), Ingress(Istio Gateway), Service Mesh, 프론트엔드, 레포 매핑, "
     "컨텍스트 간 연동(QUERY/EVENT)과 메시징 채널(Kafka). ODA 모드 전용 분기는 없습니다.",
     "ODA 산출물이 표준 diff 로 수렴했기에 기존 Plan 화면이 무분기로 동작합니다(FR-013)."),
    ("STEP 10. Impact Map — 수렴된 전술 설계", "demo-10-oda-impact-converged.png",
     "Impact Map 탭은 표준 tacticalDiff 로 수렴된 전술 설계(ProductOrder 애그리거트·ExpediteOrder "
     "명령·ProductOrderExpedited 이벤트)와 영향도를 보여줍니다. 하단 ‘샌드박스 구현 열기’가 구현 진입점입니다.",
     "여기까지가 ODA 표준 모드가 더하는 앞단이며, 이후는 기존 생애주기와 동일합니다."),
]

for h_title, img, body, note in STEPS:
    h(doc, h_title, level=2)
    doc.add_paragraph(body)
    add_screenshot(doc, img, h_title)
    if note:
        add_note(doc, note)

# 구현(Implement) 단계 — 설명 (신규 화면 없음, 기존 생애주기 재사용)
h(doc, "STEP 11. 구현(Implement) 단계", level=2)
doc.add_paragraph(
    "구현 단계는 ODA 표준 모드 전용 화면이 없습니다. Impact Map 의 ‘샌드박스 구현 열기’를 누르면 "
    "기존 Proposal 생애주기를 그대로 사용합니다 — 대상 프로젝트의 git worktree 샌드박스에서 "
    "Code 탭의 Claude Code 셀로 구현하고, 완료 후 Dual Merge(코드 머지 + 그래프 업데이트)로 반영합니다."
)
for b in [
    "표준 tacticalDiff 수렴 덕분에 샌드박스·구현·검증·수락이 분기 없이 동작합니다(SC-008).",
    "적합성 게이트는 제출 시 강제되므로, FAIL+미면제 상태면 구현까지 진행되지 않습니다.",
    "BDD .feature 산출물은 검증 계약으로 활용되며, 실제 배포·BDD 실행은 oda-componentize 영역입니다.",
]:
    doc.add_paragraph(b, style="List Bullet")
add_note(doc, "구현·검증·수락 화면은 모드와 무관하게 동일하므로 본 매뉴얼에서는 진입점까지만 다룹니다.")

# FAQ
h(doc, "자주 묻는 질문 (FAQ)")
FAQ = [
    ("ODA 표준 모드를 고르면 기존 모드는 영향을 받나요?",
     "아니요. 간소화 / 상세 DDD 모드의 동작은 그대로입니다. ODA 표준은 추가 옵션입니다."),
    ("게이트가 FAIL 인데 꼭 면제해야만 진행되나요?",
     "표준을 깨는 하드 위반은 면제(사유 기록)만이 통과 경로입니다. 위반을 해소하면 PASS 가 됩니다."),
    ("면제 사유는 어디에 남나요?",
     "적합성 리포트에 면제 사실과 사유, 시각이 기록됩니다."),
    ("표준 산출물이 다운스트림(구현)과 호환되나요?",
     "ODA 산출물은 표준 strategic/tactical diff 로 수렴되어 Impact/Tasks/구현 단계가 분기 없이 동작합니다."),
    ("ODA Component 배포까지 해주나요?",
     "아니요. 실제 배포·BDD 실행은 oda-componentize 스킬의 영역으로 본 모드 범위 밖입니다."),
]
for q, a in FAQ:
    p = doc.add_paragraph(); r = p.add_run(f"Q. {q}"); r.font.bold = True; r.font.size = Pt(10.5)
    p2 = doc.add_paragraph(f"A. {a}"); p2.paragraph_format.left_indent = Cm(0.6)
    for r in p2.runs: r.font.size = Pt(10)

doc.save(str(OUT))
print(f"WROTE {OUT} ({OUT.stat().st_size} bytes)")
