"""Render specs/044-design-import/manual.docx from the screenshots + narrative.
Run: uv run --with python-docx python scripts/make_design_import_manual_docx.py
"""
import pathlib
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SPEC = pathlib.Path("/Users/uengine/main-robo-arch/robo-architect/specs/044-design-import")
IMG = SPEC / "images"
doc = Document()

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text, italic=False, size=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return par

def img(name, width=6.3, caption=None):
    f = IMG / name
    if f.exists():
        doc.add_picture(str(f), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_heading("완성 설계 Import — 사용자 매뉴얼", 0)
p("Design 탭에서 완성된 이벤트스토밍 설계 문서를 그래프로 가져오기 · 044-design-import · 2026-06-29")

h("1. 이 기능이 하는 일", 1)
p("이미 완성된 이벤트스토밍 설계 문서(Bounded Context별로 Aggregate·Command·Event·Policy·"
  "ReadModel이 표로 정리된 마크다운 — 예: TM Forum ODA 표준 컴포넌트 캔버스)를 Design 탭에서 "
  "업로드하면, 그 설계를 LLM 재해석 없이 그대로 도메인 모델 그래프에 적재하고 Design 탭에 바로 "
  "보여줍니다. 기존 Stories 탭의 '문서 업로드'는 자연어 요구사항을 AI가 재분해하는 방식이라 완성 "
  "설계에는 맞지 않습니다. 본 기능은 충실도 보존 + Design 탭 직접 진입이 핵심입니다.")
img("01-bigpicture-oda.png", caption="ODA 표준 컴포넌트 15개 Bounded Context가 적재된 Design 탭 (좌측 목록)")

h("2. 사용 방법", 1)
h("2-1. Design 탭에서 📥(완성 설계 가져오기) 누르기", 2)
p("상단 Design 탭으로 이동한 뒤, 캔버스 좌상단 툴바의 📥 버튼을 누릅니다.")
img("03-import-modal.png", caption="완성 설계 가져오기 모달")

h("2-2. 문서 선택 + 모드 선택, 미리보기", 2)
p("완성 설계 마크다운(.md)을 고르고, 교체(기존 모델을 비우고 대체) 또는 병합(기존 모델에 추가) "
  "모드를 선택한 뒤 [미리보기]를 누릅니다. 적재될 개수와 BC별 요약, 경고, 교체 시 제거되는 기존 "
  "BC 수가 표시되며 그래프는 아직 바뀌지 않습니다.")
img("04-import-preview.png", caption="미리보기 — BC 7·Aggregate 16·Command 16·Event 23·Policy 13·사가 스파인 3, BC별 요약, 경고 15건")

h("2-3. 적재 확정 → Design 탭 렌더링", 2)
p("미리보기가 의도와 맞으면 [적재 확정]을 누릅니다. 모델이 그래프에 적재되고 Design 탭이 자동 "
  "새로고침됩니다. 좌측 Bounded Contexts 목록에서 BC를 더블클릭하면 해당 컨텍스트의 Aggregate·"
  "Command·Event·ReadModel이 이벤트스토밍 형태로 펼쳐집니다.")
img("05-design-canvas-oda.png", caption="Design 캔버스에 렌더링된 ODA 이벤트스토밍 모델 (Command·Event·Aggregate)")

h("3. 입력 문서 형식", 1)
p("이벤트스토밍 캔버스 관례를 인식합니다: 🍐 Aggregate, 🟦 Command(+actor), 🟧 Event(과거형), "
  "🟪 Policy, 🟩 Read Model. BC 섹션 = 헤딩 바로 뒤의 '| 종류 | 항목 |' 표. 컨텍스트 간 스파인 = "
  "코드 블록의 '이벤트 ─P─▶ 명령' 라인. 같은 문서를 두 번 가져오면 결과 그래프는 동일합니다(결정론).")

h("4. 주의 / 한계", 1)
p("· 본 기능은 완성된 설계(요소가 표로 정리된 문서)를 대상으로 합니다. 자유 산문 요구사항은 기존 "
  "Stories 탭 문서 업로드(AI 분해)를 쓰세요.\n"
  "· 명령→이벤트 짝이 명시되지 않은 BC는 이벤트를 첫 명령에 일괄 연결하고 경고로 알립니다.\n"
  "· 교체 모드는 기존 이벤트스토밍 모델을 비웁니다. 누적하려면 병합을 쓰세요.")

out = SPEC / "manual.docx"
doc.save(str(out))
print("saved", out)
